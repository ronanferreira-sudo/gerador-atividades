import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


def extrair_campo_avancado(texto, campo):
    """
    Extrai o valor associado a um campo numerado como:
    1. Identificação: valor
    2. Objetivos: valor
    """
    padrao = rf"\d+\.\s*{campo}\s*:?\s*(.*?)(?:\n\d+\.\s|\Z)"
    match = re.search(padrao, texto, re.DOTALL | re.IGNORECASE)
    if match:
        return match.group(1).strip()
    return ""


def extrair_subtopicos(texto, campo):
    """Extrai itens de uma seção listados com - ou •"""
    secao = extrair_campo_avancado(texto, campo)
    # Se não encontrou, tenta extrair bloco entre títulos
    if not secao:
        return []
    linhas = secao.split("\n")
    itens = []
    for linha in linhas:
        linha_strip = linha.strip()
        if linha_strip and (linha_strip.startswith("-") or linha_strip.startswith("•") or linha_strip.startswith("*")):
            itens.append(linha_strip.lstrip("-•* ").strip())
    if not itens:
        # Se não encontrou itens com marcadores, retorna o texto completo
        return [secao]
    return itens


def adicionar_cabecalho(doc, logo_path=None):
    """Adiciona cabeçalho com logo e informações"""
    cabecalho = doc.sections[0].header
    cabecalho.is_linked_to_previous = False

    if logo_path and os.path.exists(logo_path):
        try:
            paragrafo_logo = cabecalho.paragraphs[0] if cabecalho.paragraphs else cabecalho.add_paragraph()
            paragrafo_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragrafo_logo.add_run()
            run.add_picture(logo_path, width=Inches(1.5))
            paragrafo_logo.space_after = Pt(2)
        except Exception as e:
            print(f"Erro ao adicionar logo: {e}")


def adicionar_linha_horizontal(doc):
    """Adiciona uma linha horizontal"""
    paragrafo = doc.add_paragraph()
    run = paragrafo.add_run("_" * 80)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(180, 180, 180)
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return paragrafo


def gerar_docx_plano(
    titulo,
    subtitulo,
    conteudo,
    caminho_saida,
    logo_path=None
):
    doc = Document()

    # Configurar página
    secao = doc.sections[0]
    secao.top_margin = Cm(2.5)
    secao.bottom_margin = Cm(2)
    secao.left_margin = Cm(3)
    secao.right_margin = Cm(2)

    # Adicionar cabeçalho com logo se informado
    if logo_path is None:
        # Tenta encontrar o logo em locais comuns
        for possivel in [
            "static/img/logo.png",
            "static/img/logo.jpg",
            os.path.join(os.path.dirname(os.path.dirname(__file__)), "static/img/logo.png"),
            os.path.join(os.path.dirname(os.path.dirname(__file__)), "static/img/logo.jpg"),
        ]:
            if os.path.exists(possivel):
                logo_path = possivel
                break

    # Cabeçalho com logo
    if logo_path and os.path.exists(logo_path):
        try:
            cabecalho = doc.sections[0].header
            cabecalho.is_linked_to_previous = False

            # Apenas a logo, sem texto
            paragrafo_logo = cabecalho.paragraphs[0] if cabecalho.paragraphs else cabecalho.add_paragraph()
            paragrafo_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_logo = paragrafo_logo.add_run()
            run_logo.add_picture(logo_path, width=Inches(1.2))
            paragrafo_logo.space_after = Pt(2)
        except Exception as e:
            print(f"Erro ao adicionar logo no cabeçalho: {e}")

    # Título = nome do curso (subtitulo), alinhado à esquerda
    if subtitulo:
        p_titulo = doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_titulo.add_run(subtitulo)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 70, 150)
        p_titulo.space_after = Pt(10)

    # Processar conteúdo - igual ao template HTML
    linhas = conteudo.split("\n")
    for linha in linhas:
        linha_strip = linha.strip()

        # Identificar seções principais (numeradas ou por título)
        if re.match(r'^\d+[\.\:\-–]\s*(Identificação|Objetivos|Conteudo|Conteúdo|Estrategia|Estratégia|Recursos|Avaliação|Avaliacao|Metodologia|Titulo|Título|Introdução|Introducao|Desenvolvimento|Conclusão|Conclusao)', linha_strip, re.IGNORECASE) or \
           re.match(r'^(Identificação|Objetivos|Conteudo Programático|Conteúdo Programático|Estrategia Didática|Estratégia Didática|Recursos Didáticos|Recursos Didaticos|Avaliação|Avaliacao|Metodologia|Introdução|Introducao)', linha_strip, re.IGNORECASE) or \
           re.match(r'^DIA\s+\d+', linha_strip, re.IGNORECASE) or \
           re.match(r'^PLANO', linha_strip, re.IGNORECASE):

            # Seção principal - heading
            p = doc.add_paragraph()
            p.space_before = Pt(14)
            p.space_after = Pt(4)
            run = p.add_run(linha_strip)
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0, 70, 150)

        elif linha_strip == '':
            # Linha em branco
            doc.add_paragraph().space_after = Pt(2)

        elif linha_strip.startswith('-') or linha_strip.startswith('•') or linha_strip.startswith('*'):
            # Item de lista
            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            run = p.add_run(linha_strip.lstrip("-•* ").strip())
            run.font.size = Pt(11)
            p.space_before = Pt(0)
            p.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(1)

        else:
            # Parágrafo normal
            p = doc.add_paragraph()
            run = p.add_run(linha_strip)
            run.font.size = Pt(11)
            p.space_before = Pt(1)
            p.space_after = Pt(1)

    # Salvar
    pasta = os.path.dirname(caminho_saida)
    if pasta:
        os.makedirs(pasta, exist_ok=True)

    doc.save(caminho_saida)
    print(f"✅ WORD GERADO COM SUCESSO: {caminho_saida}")