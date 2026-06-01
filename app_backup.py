from flask import Flask, render_template, request, redirect, session
from docx import Document
from flask import send_file
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import os
from flask import Flask, render_template, request, redirect
from database.conexao import conexao, cursor
from ia.gerador import gerar_atividade

app = Flask(__name__)
app.secret_key = 'gerador_atividades_2026'


@app.route('/', methods=['GET', 'POST'])
def index():

    if request.method == 'POST':

        curso = request.form['curso']
        disciplina = request.form['disciplina']
        conteudo = request.form['conteudo']
        dificuldade = request.form['dificuldade']

        print("Gerando atividade com IA...")

        atividade_gerada = gerar_atividade(
            conteudo,
            dificuldade
        )

        cursor.execute("""
            INSERT INTO atividades (
                curso,
                disciplina,
                conteudo,
                dificuldade,
                atividade_gerada
            )
            VALUES (%s, %s, %s, %s, %s)
        """, (
            curso,
            disciplina,
            conteudo,
            dificuldade,
            atividade_gerada
        ))

        conexao.commit()

        return redirect('/atividades')

    return render_template('index.html')


@app.route('/atividades')
def atividades():

    cursor.execute("""
        SELECT *
        FROM atividades
        ORDER BY id DESC
    """)

    dados = cursor.fetchall()

    return render_template(
        'atividades.html',
        atividades=dados
    )


@app.route('/deletar/<int:id>')
def deletar(id):

    cursor.execute(
        "DELETE FROM atividades WHERE id = %s",
        (id,)
    )

    conexao.commit()

    return redirect('/atividades')


@app.route('/editar/<int:id>', methods=['GET', 'POST'])
def editar(id):

    if request.method == 'POST':

        curso = request.form['curso']
        disciplina = request.form['disciplina']
        conteudo = request.form['conteudo']
        dificuldade = request.form['dificuldade']
        atividade_gerada = request.form['atividade_gerada']

        cursor.execute("""
            UPDATE atividades
            SET curso=%s,
                disciplina=%s,
                conteudo=%s,
                dificuldade=%s,
                atividade_gerada=%s
            WHERE id=%s
        """, (
            curso,
            disciplina,
            conteudo,
            dificuldade,
            atividade_gerada,
            id
        ))

        conexao.commit()

        return redirect('/atividades')

    cursor.execute(
        "SELECT * FROM atividades WHERE id=%s",
        (id,)
    )

    atividade = cursor.fetchone()

    return render_template(
        'editar.html',
        atividade=atividade
    )
@app.route('/pdf/<int:id>')
def gerar_pdf(id):

    cursor.execute(
        "SELECT * FROM atividades WHERE id=%s",
        (id,)
    )

    atividade = cursor.fetchone()

    arquivo = f"atividade_{id}.pdf"

    doc = SimpleDocTemplate(arquivo)

    estilos = getSampleStyleSheet()

    elementos = []

    elementos.append(
        Paragraph(
            f"<b>Curso:</b> {atividade[1]}",
            estilos['Normal']
        )
    )

    elementos.append(
        Paragraph(
            f"<b>Disciplina:</b> {atividade[2]}",
            estilos['Normal']
        )
    )

    elementos.append(
        Paragraph(
            f"<b>Conteúdo:</b> {atividade[3]}",
            estilos['Normal']
        )
    )

    elementos.append(
        Paragraph(
            f"<b>Dificuldade:</b> {atividade[4]}",
            estilos['Normal']
        )
    )

    elementos.append(
        Paragraph(
            atividade[5].replace("\n", "<br/>"),
            estilos['Normal']
        )
    )

    doc.build(elementos)

    return send_file(
        arquivo,
        as_attachment=True
    )
@app.route('/word/<int:id>')
def gerar_word(id):

    cursor.execute(
        "SELECT * FROM atividades WHERE id=%s",
        (id,)
    )

    atividade = cursor.fetchone()

    documento = Document()

    documento.add_heading(
        'Atividade Gerada pela IA',
        level=1
    )

    documento.add_paragraph(
        f'Curso: {atividade[1]}'
    )

    documento.add_paragraph(
        f'Disciplina: {atividade[2]}'
    )

    documento.add_paragraph(
        f'Conteúdo: {atividade[3]}'
    )

    documento.add_paragraph(
        f'Dificuldade: {atividade[4]}'
    )

    documento.add_paragraph('')

    documento.add_paragraph(
        atividade[5]
    )

    arquivo = f'atividade_{id}.docx'

    documento.save(arquivo)

    return send_file(
        arquivo,
        as_attachment=True
    )
@app.route('/login', methods=['GET', 'POST'])
def login():

    if request.method == 'POST':

        email = request.form['email']
        senha = request.form['senha']

        cursor.execute("""
            SELECT id, nome, email, perfil
            FROM usuarios
            WHERE email=%s AND senha=%s
        """, (email, senha))

        usuario = cursor.fetchone()

        if usuario:

            session['usuario_id'] = usuario[0]
            session['nome'] = usuario[1]
            session['perfil'] = usuario[3]

            return redirect('/')

        return "Login inválido"

    return render_template('login.html')

if __name__ == '__main__':
    app.run(debug=True)