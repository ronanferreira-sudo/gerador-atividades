from PyPDF2 import PdfReader
from werkzeug.utils import secure_filename
from ia.plano_aula import gerar_plano_aula
from flask import Flask, render_template, request, redirect, session, send_file
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import os
from ia.gerador import gerar_atividade
from database.conexao import conexao, cursor

app = Flask(__name__)
app.secret_key = "gerador_atividades_2026"


def pode_acessar_atividade(id):

    if session["perfil"] == "admin":
        return True

    cursor.execute(
        """
        SELECT usuario_id
        FROM atividades
        WHERE id = %s
    """,
        (id,),
    )

    atividade = cursor.fetchone()

    if not atividade:
        return False

    return atividade[0] == session["usuario_id"]


# =========================
# LOGIN
# =========================
@app.route("/cadastro", methods=["GET", "POST"])
def cadastro():

    if request.method == "POST":

        nome = request.form["nome"]
        email = request.form["email"]
        senha = request.form["senha"]

        cursor.execute(
            """
            INSERT INTO usuarios (
                nome,
                email,
                senha,
                perfil
            )
            VALUES (%s,%s,%s,%s)
        """,
            (nome, email, senha, "professor"),
        )

        conexao.commit()

        return redirect("/login")

    return render_template("cadastro.html")


@app.route("/login", methods=["GET", "POST"])
def login():

    if request.method == "POST":

        email = request.form["email"]
        senha = request.form["senha"]

        cursor.execute(
            """
            SELECT id, nome, email, perfil
            FROM usuarios
            WHERE email=%s AND senha=%s
        """,
            (email, senha),
        )

        usuario = cursor.fetchone()

        if usuario:
            session["usuario_id"] = usuario[0]
            session["nome"] = usuario[1]
            session["perfil"] = usuario[3]

            return redirect("/dashboard")

        return "Login inválido"

    return render_template("login.html")


@app.route("/logout")
def logout():
    session.clear()
    return redirect("/login")


# =========================
# INDEX (GERAR IA)
# =========================
@app.route("/", methods=["GET", "POST"])
def index():

    if "usuario_id" not in session:
        return redirect("/login")

    if request.method == "POST":

        curso = request.form["curso"]
        disciplina = request.form["disciplina"]
        conteudo = request.form["conteudo"]
        dificuldade = request.form["dificuldade"]

        tipo = request.form["tipo"]
        quantidade = request.form["quantidade"]

        print("Gerando atividade com IA...")

        atividade_gerada = gerar_atividade(conteudo, dificuldade, tipo, quantidade)

        usuario_id = session["usuario_id"]

        cursor.execute(
            """
            INSERT INTO atividades (
                curso,
                disciplina,
                conteudo,
                dificuldade,
                atividade_gerada,
                usuario_id
            )
            VALUES (%s,%s,%s,%s,%s,%s)
        """,
            (curso, disciplina, conteudo, dificuldade, atividade_gerada, usuario_id),
        )

        conexao.commit()

        return redirect("/atividades")

    return render_template("index.html")


# =========================
# LISTAR + PESQUISA
# =========================
@app.route("/atividades")
def atividades():

    if "usuario_id" not in session:
        return redirect("/login")

    busca = request.args.get("busca")

    if session["perfil"] == "admin":

        if busca:
            cursor.execute(
                """
                SELECT *
                FROM atividades
                WHERE curso ILIKE %s
                   OR disciplina ILIKE %s
                   OR conteudo ILIKE %s
                ORDER BY id DESC
            """,
                (f"%{busca}%", f"%{busca}%", f"%{busca}%"),
            )
        else:
            cursor.execute("""
                SELECT *
                FROM atividades
                ORDER BY id DESC
            """)

    else:

        if busca:
            cursor.execute(
                """
                SELECT *
                FROM atividades
                WHERE usuario_id = %s
                  AND (curso ILIKE %s
                   OR disciplina ILIKE %s
                   OR conteudo ILIKE %s)
                ORDER BY id DESC
            """,
                (session["usuario_id"], f"%{busca}%", f"%{busca}%", f"%{busca}%"),
            )
        else:
            cursor.execute(
                """
                SELECT *
                FROM atividades
                WHERE usuario_id = %s
                ORDER BY id DESC
            """,
                (session["usuario_id"],),
            )

    dados = cursor.fetchall()

    return render_template("atividades.html", atividades=dados)


# =========================
# DELETE
# =========================
@app.route("/deletar/<int:id>")
def deletar(id):

    if not pode_acessar_atividade(id):
        return "Acesso negado"

    cursor.execute("DELETE FROM atividades WHERE id=%s", (id,))

    conexao.commit()

    return redirect("/atividades")


# =========================
# EDITAR
# =========================
@app.route("/editar/<int:id>", methods=["GET", "POST"])
def editar(id):

    if not pode_acessar_atividade(id):
        return "Acesso negado"

    if request.method == "POST":

        curso = request.form["curso"]
        disciplina = request.form["disciplina"]
        conteudo = request.form["conteudo"]
        dificuldade = request.form["dificuldade"]
        atividade_gerada = request.form["atividade_gerada"]

        cursor.execute(
            """
            UPDATE atividades
            SET curso=%s,
                disciplina=%s,
                conteudo=%s,
                dificuldade=%s,
                atividade_gerada=%s
            WHERE id=%s
        """,
            (curso, disciplina, conteudo, dificuldade, atividade_gerada, id),
        )

        conexao.commit()

        return redirect("/atividades")

    cursor.execute("SELECT * FROM atividades WHERE id=%s", (id,))
    atividade = cursor.fetchone()

    if not atividade:
        return "Atividade não encontrada"

    return render_template("editar.html", atividade=atividade)


# =========================
# PDF
# =========================
@app.route("/pdf/<int:id>")
def gerar_pdf(id):
    if not pode_acessar_atividade(id):
        return "Acesso negado"

    cursor.execute("SELECT * FROM atividades WHERE id=%s", (id,))
    atividade = cursor.fetchone()

    arquivo = f"atividade_{id}.pdf"
    doc = SimpleDocTemplate(arquivo)
    estilos = getSampleStyleSheet()

    elementos = []

    elementos.append(Paragraph(f"<b>Curso:</b> {atividade[1]}", estilos["Normal"]))
    elementos.append(Paragraph(f"<b>Disciplina:</b> {atividade[2]}", estilos["Normal"]))
    elementos.append(Paragraph(f"<b>Conteúdo:</b> {atividade[3]}", estilos["Normal"]))
    elementos.append(
        Paragraph(f"<b>Dificuldade:</b> {atividade[4]}", estilos["Normal"])
    )
    elementos.append(Paragraph(atividade[5].replace("\n", "<br/>"), estilos["Normal"]))

    doc.build(elementos)

    return send_file(arquivo, as_attachment=True)


# =========================
# WORD
# =========================
@app.route("/word/<int:id>")
def gerar_word(id):
    if not pode_acessar_atividade(id):
        return "Acesso negado"

    cursor.execute("SELECT * FROM atividades WHERE id=%s", (id,))
    atividade = cursor.fetchone()

    doc = Document()

    doc.add_heading("Atividade Gerada pela IA", 1)

    doc.add_paragraph(f"Curso: {atividade[1]}")
    doc.add_paragraph(f"Disciplina: {atividade[2]}")
    doc.add_paragraph(f"Conteúdo: {atividade[3]}")
    doc.add_paragraph(f"Dificuldade: {atividade[4]}")
    doc.add_paragraph("")
    doc.add_paragraph(atividade[5])

    arquivo = f"atividade_{id}.docx"
    doc.save(arquivo)

    return send_file(arquivo, as_attachment=True)


# =========================
# REGENERAR IA
# =========================
@app.route("/regenerar/<int:id>")
def regenerar(id):
    if not pode_acessar_atividade(id):
        return "Acesso negado"

    cursor.execute(
        """
        SELECT conteudo, dificuldade
        FROM atividades
        WHERE id=%s
    """,
        (id,),
    )

    atividade = cursor.fetchone()

    nova = gerar_atividade(atividade[0], atividade[1])

    cursor.execute(
        """
        UPDATE atividades
        SET atividade_gerada=%s
        WHERE id=%s
    """,
        (nova, id),
    )

    conexao.commit()

    return redirect("/atividades")


# =========================
# DASHBOARD
# =========================
@app.route("/dashboard")
def dashboard():

    cursor.execute("""
        SELECT
            curso,
            COUNT(*)
        FROM atividades
        GROUP BY curso
        ORDER BY curso
    """)

    cursos = cursor.fetchall()

    return render_template("dashboard.html", cursos=cursos)


# =========================
# PLANOS DE AULA
# =========================
@app.route("/planos", methods=["GET", "POST"])
def planos():

    if "usuario_id" not in session:
        return redirect("/login")

    if request.method == "POST":

        nome_curso = request.form["nome_curso"]
        carga_horaria = int(request.form["carga_horaria"])
        aulas_por_dia = int(request.form["aulas_por_dia"])

        arquivo = request.files["arquivo_pdf"]

        nome_arquivo = secure_filename(arquivo.filename)

        pasta_upload = "uploads"

        if not os.path.exists(pasta_upload):
            os.makedirs(pasta_upload)

        caminho = os.path.join(pasta_upload, nome_arquivo)

        arquivo.save(caminho)

        print("PDF salvo!")

        # =========================
        # EXTRAIR TEXTO DO PDF
        # =========================

        reader = PdfReader(caminho)

        texto = ""

        for pagina in reader.pages:

            conteudo = pagina.extract_text()

            if conteudo:
                texto += conteudo + "\n"

        print("Texto extraído!")
        print("Quantidade de caracteres:", len(texto))

        # =========================
        # SALVAR CURSO
        # =========================

        cursor.execute(
            """
            INSERT INTO cursos_plano (
                nome_curso,
                carga_horaria,
                aulas_por_dia,
                usuario_id
            )
            VALUES (%s,%s,%s,%s)
            RETURNING id
        """,
            (nome_curso, carga_horaria, aulas_por_dia, session["usuario_id"]),
        )

        curso_id = cursor.fetchone()[0]

        print("Curso salvo!")

        # =========================
        # GERAR PLANO COM IA
        # =========================

        print("Enviando para IA...")

        plano_gerado = gerar_plano_aula(texto[:1000], carga_horaria, aulas_por_dia)

        print("IA respondeu!")

        # =========================
        # SEPARAR DIAS
        # =========================

        planos = plano_gerado.split("### DIA")

        print("Planos encontrados:", len(planos))

        # =========================
        # SALVAR CADA DIA
        # =========================

        for plano in planos:

            plano = plano.strip()

            if not plano:
                continue

            linhas = plano.split("\n")

            try:
                dia = int(linhas[0].strip())
            except:
                dia = 0

            cursor.execute(
                """
                INSERT INTO planos_aula (
                    curso_id,
                    dia,
                    tema,
                    plano_gerado,
                    usuario_id
                )
                VALUES (%s,%s,%s,%s,%s)
            """,
                (curso_id, dia, f"Plano Dia {dia}", plano, session["usuario_id"]),
            )

        conexao.commit()

        print("Planos salvos no banco!")

        return redirect("/listar_cursos")

    return render_template("planos.html")


# =========================
# LISTAR PLANOS
# =========================
@app.route("/listar_planos")
def listar_planos():

    if "usuario_id" not in session:
        return redirect("/login")

    if session["perfil"] == "admin":

        cursor.execute("""
            SELECT *
            FROM planos_aula
            ORDER BY id DESC
        """)

    else:

        cursor.execute(
            """
            SELECT *
            FROM planos_aula
            WHERE usuario_id = %s
            ORDER BY id DESC
        """,
            (session["usuario_id"],),
        )

    planos = cursor.fetchall()

    return render_template("listar_planos.html", planos=planos)


# =========================
# VISUALIZAR PLANOS
# =========================


@app.route("/visualizar_plano/<int:id>")
def visualizar_plano(id):

    cursor.execute(
        """
        SELECT *
        FROM planos_aula
        WHERE id = %s
    """,
        (id,),
    )

    plano = cursor.fetchone()

    return render_template("visualizar_plano.html", plano=plano)


# =========================
# LISTAR CURSOS
# =========================
@app.route("/listar_cursos")
def listar_cursos():

    if "usuario_id" not in session:
        return redirect("/login")

    if session["perfil"] == "admin":

        cursor.execute("""
            SELECT *
            FROM cursos_plano
            ORDER BY id DESC
        """)

    else:

        cursor.execute(
            """
            SELECT *
            FROM cursos_plano
            WHERE usuario_id = %s
            ORDER BY id DESC
        """,
            (session["usuario_id"],),
        )

    cursos = cursor.fetchall()

    return render_template("listar_cursos.html", cursos=cursos)


# =========================
# VISUALIZAR CURSO
# =========================
@app.route("/curso/<int:curso_id>")
def visualizar_curso(curso_id):

    if "usuario_id" not in session:
        return redirect("/login")

    cursor.execute(
        """
        SELECT *
        FROM planos_aula
        WHERE curso_id = %s
        ORDER BY dia
    """,
        (curso_id,),
    )

    planos = cursor.fetchall()

    return render_template("planos_curso.html", planos=planos, curso_id=curso_id)


# =========================
# START
# =========================
if __name__ == "__main__":
    app.run(debug=True)
